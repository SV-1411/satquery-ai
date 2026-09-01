from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r'D:\satquery-ai\docs\SatQuery_AI_Operator_and_Implementation_Handbook.docx')
BLUE, NAVY, LIGHT = '2E74B5', '1F4D78', 'E8EEF5'

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)

def margins(cell):
    tcPr = cell._tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for side in ('top', 'start', 'bottom', 'end'):
        n = tcMar.find(qn(f'w:{side}'))
        if n is None: n = OxmlElement(f'w:{side}'); tcMar.append(n)
        n.set(qn('w:w'), '100' if side in ('start','end') else '80'); n.set(qn('w:type'), 'dxa')

def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'; t.autofit = False; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; shade(c, LIGHT); margins(c); c.width = Pt(widths[i]/20)
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0); r = p.add_run(h); r.bold = True; r.font.color.rgb = RGBColor.from_string(NAVY); r.font.size = Pt(9)
    for row in rows:
        cs = t.add_row().cells
        for i, value in enumerate(row):
            c = cs[i]; margins(c); c.width = Pt(widths[i]/20); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.05; r = p.add_run(str(value)); r.font.size = Pt(9)
    return t

def p(doc, text='', style=None, after=6, align=None):
    x = doc.add_paragraph(style=style); x.paragraph_format.space_after = Pt(after)
    if align: x.alignment = align
    if text: x.add_run(text)
    return x

def bullet(doc, text, level=0):
    x = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2'); x.add_run(text); return x

def numbered(doc, text):
    x = doc.add_paragraph(style='List Number'); x.add_run(text); return x

def code(doc, text):
    x = doc.add_paragraph(); x.paragraph_format.left_indent = Inches(.18); x.paragraph_format.space_after = Pt(8); x.paragraph_format.line_spacing = 1.0
    r = x.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(28,39,51); return x

doc = Document(); sec = doc.sections[0]
sec.top_margin = Inches(.7); sec.bottom_margin = Inches(.65); sec.left_margin = Inches(.75); sec.right_margin = Inches(.75)
normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(11); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.15
for name, size, color, before, after in [('Heading 1',16,BLUE,18,10), ('Heading 2',13,BLUE,14,7), ('Heading 3',12,NAVY,10,5)]:
    s = doc.styles[name]; s.font.name = 'Calibri'; s.font.size = Pt(size); s.font.bold = True; s.font.color.rgb = RGBColor.from_string(color); s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after); s.paragraph_format.keep_with_next = True
hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT; rr = hp.add_run('SATQUERY AI  /  SENTRY  •  OPERATOR HANDBOOK'); rr.bold=True; rr.font.color.rgb=RGBColor.from_string(BLUE); rr.font.size=Pt(8)
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER; rr = fp.add_run('SatQuery AI prototype 0.1  •  Evidence is a claim with a location, a source and a limit.'); rr.font.color.rgb=RGBColor(102,119,136); rr.font.size=Pt(8)

# Cover
x = p(doc, 'SATQUERY AI  /  SENTRY', after=3); x.runs[0].bold=True; x.runs[0].font.color.rgb=RGBColor.from_string(BLUE); x.runs[0].font.size=Pt(11)
x = p(doc, 'Earth, under evidence.', after=8); x.runs[0].bold=True; x.runs[0].font.size=Pt(30)
x = p(doc, 'Operator, test and implementation handbook', after=22); x.runs[0].font.size=Pt(15); x.runs[0].font.color.rgb=RGBColor.from_string(BLUE)
table(doc, ['Document','Value'], [('Version','Prototype 0.1 / September 2026'),('Audience','Hackathon judges, operators, developers and future customers'),('Scope','Single-image VQA + grounding/captioning, bi-temporal change reasoning, optical–SAR fusion'),('Current deployment','Owner-only Sites deployment at satquery-sentry.shrutikaverma22.chatgpt.site')], [1800,7560])
p(doc, 'This handbook is deliberately practical. It explains the product idea, scientific design, current prototype behaviour, exact interface actions, repeatable tests, deployment, and the path to a production-grade remote-sensing platform.', after=8)
doc.add_page_break()

doc.add_heading('1. Executive summary', 1)
p(doc, 'SatQuery AI is an evidence-first assistant for asking plain-language questions about satellite observations. A user supplies one image, a co-registered optical–SAR pair, or two dates of the same area, then asks what matters in ordinary language. An agentic controller validates the inputs, classifies the requested task, selects a specialist remote-sensing tool, and returns an evidence contract: what the system believes, where it occurs, how large it may be, which sensor supports it, what could be wrong, and how confident the system is.')
p(doc, 'The product is not a chat wrapper around a general model. Its defensible idea is a decision layer around specialised Earth-observation models. The interface makes the chain of evidence visible, allows abstention when sensors disagree, and preserves an execution trace that can be audited by an analyst, emergency team or infrastructure operator.')
doc.add_heading('What makes it distinctive', 2)
for t in ['Evidence Contract: every answer contains a claim, location, magnitude, sensor case, confidence and next observation/limit.','Sensor Contribution Ledger: optical and SAR evidence are reported separately before fusion, so users can see whether a conclusion is genuinely cross-modal.','Change Ledger: bi-temporal reasoning records before/after observations, changed region, direction and uncertainty instead of only a class label.','Calibrated abstention: the assistant can withhold a region or magnitude when cloud, registration error or sensor disagreement makes the answer unsafe.','Query-driven routing: the task is selected from the query and input configuration; the workflow is observable instead of hidden behind one generic VLM.']: bullet(doc,t)
doc.add_heading('Prototype boundary', 2)
p(doc, 'The current web prototype performs real multipart validation and raster-header inspection, but its returned claims are deterministic representative outputs. It is not yet running a fine-tuned VQA, grounding, change model or optical–SAR fusion network. This clean vertical slice keeps the contract and orchestration testable while the roadmap below replaces each representative tool with benchmarked models.')

doc.add_heading('2. Problem statement and required scope', 1)
p(doc, 'Remote-sensing systems are often isolated applications: one classifier, one detector, one VQA model or one change detector. That forces non-specialists to understand sensors, GIS conventions, model parameters and expected input shapes. SatQuery AI treats the question and the available observations as the starting point.')
table(doc, ['Requirement','SatQuery implementation'], [('Remote-sensing adaptation','Use BigEarthNet.txt (Sentinel-1 SAR + Sentinel-2 multispectral + text) for image–text adaptation; the model registry preserves sensor tokens.'),('Single-image VQA','SINGLE_IMAGE_VQA for land cover, objects, roads, water and built-up questions.'),('Second single-image task','TEXT_GUIDED_GROUNDING for highlight/locate/where/mark/region prompts; captioning can be added beside it.'),('Bi-temporal analysis','BI_TEMPORAL_CHANGE_VQA when two corresponding observations and a change/increase/decrease query are present.'),('Optical–SAR analysis','OPTICAL_SAR_FUSION for co-registered pairs and prompts mentioning SAR, radar, optical, together or fuse.'),('Agentic orchestration','Input inspector → query router → specialist registry → evidence builder → trace/report.'),('Accepted inputs','GeoTIFF/TIFF for geospatial imagery; PNG/JPEG for prescribed benchmark demonstrations.')], [2200,7160])
doc.add_heading('Datasets and evaluation intent', 2)
p(doc, 'BigEarthNet.txt is the primary adaptation source named by the problem. VRSBench supports captioning, grounding and VQA; RSVQA supports single-image VQA; CDVQA supports change-based VQA. The closed ISRO/SAC set is expected to contain co-registered Cartosat-2S optical and RISAT SAR pairs with answers, boxes or masks. Keep train, validation and public-test handling separate and document the licence and provenance of every additional model or dataset.')

doc.add_heading('3. Product workflow and architecture', 1)
table(doc, ['Stage','What happens','Observable output'], [('1. Inspect','Check count, extension, byte size, raster signature, dimensions and band estimate.','Input summary with PASSED/PARTIAL/REJECTED.'),('2. Interpret','Classify the natural-language query with image count and selected case.','Task stamp such as BI_TEMPORAL_CHANGE_VQA.'),('3. Select','Resolve task against a registry of specialists and allowed parameters.','Trace names the selected tool.'),('4. Execute','Run VQA, grounding, change or fusion inference; align outputs to the input grid.','Text claim plus optional map, mask or box.'),('5. Validate','Compare modalities, check confidence and enforce abstention gates.','Sensor case, limit and decision.'),('6. Report','Render the same contract in UI and download it.','Evidence report and trace.')], [1200,5000,3160])
code(doc, 'Browser UI\n    │  query + 0/1/2 observations\n    ▼\nPOST /api/analyse (multipart/form-data)\n    │\n    ├─ Input Inspector: extension, size, signature, dimensions, bands\n    ├─ Query Router: VQA | Grounding | Change VQA | Optical–SAR Fusion\n    ├─ Specialist Registry: model + permitted parameters\n    ├─ Evidence Builder: claim, where, magnitude, sensor case, limit\n    └─ Trace + report\n    ▼\nEvidence map + confidence + auditable execution summary')
p(doc, 'The governing principle is simple: an answer is incomplete until a human can tell what was observed, where it was observed, which source contributed and what would change the conclusion.')

doc.add_heading('4. Using the deployed interface', 1)
p(doc, 'Open https://satquery-sentry.shrutikaverma22.chatgpt.site. The controls and typography intentionally use a black-and-white brutalist system. The evidence field uses a higher-quality full-colour satellite-style basemap so spatial evidence is easier to read.')
doc.add_heading('Screen map', 2)
for t in ['01 / OBSERVATIONS: choose a curated case or load your own files. Case cards communicate expected sensor and temporal configuration.','02 / EVIDENCE FIELD: switch FUSED, OPTICAL or SAR to inspect the evidence through a modality lens.','ASK THE OBSERVATION: type a question, choose a suggestion, and press RUN →. The button reads RUNNING… while processing.','03 / EVIDENCE CONTRACT: read claim, location, magnitude, sensor support, limitation, confidence and decision. Expand OBSERVABLE EXECUTION TRACE for routing steps.','DOWNLOAD EVIDENCE REPORT ↓: save the displayed contract and trace as a plain-text report for a ticket or analyst hand-off.']: bullet(doc,t)
doc.add_heading('Curated cases', 2)
table(doc, ['Case','Intended demonstration','Expected routing'], [('CASE_001 / ASSAM FLOOD','Sentinel-2 optical + Sentinel-1 SAR, two dates. Ask what changed and which sensor supports it.','BI_TEMPORAL_CHANGE_VQA'),('CASE_002 / URBAN WATER','RISAT SAR single-image question. Ask what is visible or where water is located.','SINGLE_IMAGE_VQA or TEXT_GUIDED_GROUNDING'),('CASE_003 / BUILT CHANGE','Cartosat-2S optical, two dates. Ask whether built-up area increased, decreased or stayed unchanged.','BI_TEMPORAL_CHANGE_VQA')], [2100,4860,2400])
doc.add_heading('What each control does', 2)
for t in ['Click a case card. The active card is selected and existing files are cleared so configurations cannot be mixed accidentally.','Click + ADD OBSERVATION. Select up to two .tif, .tiff, .png, .jpg or .jpeg files. File names and header-derived metadata appear in the left panel.','Choose FUSED, OPTICAL or SAR. This changes the evidence view; it does not invent a missing modality.','Edit the question or click a suggestion. Prefer specific language such as “What changed between these dates?” or “Highlight water-covered regions”.','Press RUN →. The browser sends a multipart request to /api/analyse and replaces the contract, confidence and trace when it returns.','Press DOWNLOAD EVIDENCE REPORT ↓. A file named satquery-<case>-report.txt is generated locally by the browser.']: numbered(doc,t)

doc.add_heading('5. Test recipes and expected results', 1)
table(doc, ['Test','Action','Expected result'], [('Default change','Keep CASE_001; run “What changed, where, and which sensor supports it?”','BI_TEMPORAL_CHANGE_VQA; confidence 82; newly inundated land likely; Zones A + B; CHANGE_ANALYST trace.'),('Single-image VQA','CASE_002; run “Describe the land-cover and major objects visible in this image.”','SINGLE_IMAGE_VQA; confidence 76; water, built-up land and road corridors.'),('Grounding','Run “Highlight water-covered regions”.','TEXT_GUIDED_GROUNDING; map regions A and B remain visible when evidence is sufficient.'),('Optical–SAR fusion','With two observations, run “Use the optical and SAR images together to identify built-up and water-covered regions.”','OPTICAL_SAR_FUSION; sensor case reports SAR strong support and optical partial agreement.'),('Abstention','Add “cloud” or “uncertain”, for example “What changed despite cloud uncertainty?”','Confidence 43; INSUFFICIENT EVIDENCE; regions withheld; next observation recommends cloud-free optical or analyst review.'),('PNG metadata','Upload D:\\satquery-ai\\public\\og.png.','PNG header parsed; 1200 × 630, estimated 3 bands, PASSED.'),('Unsupported file','Try a PDF.','UI input error; API responds HTTP 415 with filename and supported formats.')], [1700,4300,3360])
doc.add_heading('Direct API smoke test', 2)
code(doc, 'Set-Location D:\\satquery-ai\n$body = @{\n  caseId = "assam-flood"\n  query = "What changed, where, and which sensor supports it?"\n  files = Get-Item .\\public\\og.png\n}\nInvoke-RestMethod -Uri http://localhost:3000/api/analyse -Method Post -Form $body')
p(doc, 'The response is JSON. Check task, confidence, inputSummary and trace first, then claim/where/magnitude/limit. A request with no files demonstrates a curated case; an upload exercises byte-level header inspection.')

doc.add_heading('6. Local development', 1)
for t in ['Install Node.js 22 or newer and Git.','In PowerShell: Set-Location D:\\satquery-ai','Install dependencies if needed: npm install','Start: npm run dev','Open http://localhost:3000','Before committing, run npm run lint and npm run build.']: numbered(doc,t)
doc.add_heading('Useful files', 2)
table(doc, ['Path','Purpose'], [('app/satquery-console.tsx','Client UI, cases, modes, upload handling, request submission and report download.'),('app/api/analyse/route.ts','Multipart API, raster parsers, task classifier and current representative contract.'),('app/globals.css','Brutalist control system, map presentation and modality overlays.'),('public/evidence-map.png','Full-colour satellite-style evidence basemap.'),('public/og.png','Social preview image.'),('app/layout.tsx','Document metadata and social-card references.'),('.openai/hosting.json','Sites project configuration.')], [2800,6560])

doc.add_heading('7. API contract', 1)
doc.add_heading('Request', 2)
p(doc, 'POST /api/analyse with multipart/form-data. Fields: caseId (optional), query (required, at least five characters), and zero to two files under files. Accepted extensions are .tif/.tiff/.png/.jpg/.jpeg; the prototype rejects files over 50 MB.')
doc.add_heading('Response fields', 2)
table(doc, ['Field','Meaning'], [('task','Selected specialist workflow identifier.'),('claim','Plain-language conclusion released by the evidence builder.'),('where','Named or spatially grounded area associated with the claim.'),('magnitude','Estimated area/quantity and uncertainty.'),('sensorCase','Per-sensor support or disagreement summary.'),('limit','Known limitation and recommended next observation.'),('confidence / decision','Calibrated score and operational release state.'),('inputSummary','File name, format, bytes, dimensions/bands and status.'),('trace','Ordered step/tool/status rows making orchestration auditable.')], [2300,7060])
doc.add_heading('Error semantics', 2)
table(doc, ['HTTP','Condition','Action'], [('400','Body is not multipart form data.','Send FormData; do not JSON-encode file upload.'),('415','Unsupported extension, invalid signature or oversized file.','Use a real GeoTIFF/TIFF/PNG/JPEG under 50 MB.'),('422','Question missing or too vague.','Write a concrete question of at least five characters.')], [1200,4200,3960])

doc.add_heading('8. Scientific design for production', 1)
p(doc, 'Retain the current contract while replacing representative outputs with measured models. Model selection should be a constrained classification problem, not an invitation for an LLM to invent tools or parameters.')
doc.add_heading('Adaptation and registry', 2)
for t in ['Use BigEarthNet.txt to adapt a dual-encoder or encoder–decoder image–text representation across Sentinel-1 SAR, Sentinel-2 multispectral imagery and text. Preserve sensor tokens and acquisition metadata.','Register models with task, modality, required image count, alignment requirement, supported bands, version, checksum, calibration set and permitted parameters.','Keep frozen validation slices per task for calibration. A high raw logit is not an operational confidence score.']: bullet(doc,t)
doc.add_heading('Fusion, change and grounding', 2)
p(doc, 'For optical–SAR pairs, align grids and footprints first, then encode each modality separately. Fuse only after recording per-modality evidence. A disagreement head should detect cloud, layover, shadow, speckle and registration errors. For bi-temporal reasoning, predict change probability, direction and semantic description; expose a mask only after registration, morphology and confidence gates. For grounding, return boxes/polygons and, for GeoTIFF, convert them to the source CRS. Store geometry, CRS and score in the report.')

doc.add_heading('9. Roadmap from prototype to startup', 1)
table(doc, ['Phase','Build','Exit criteria'], [('0 / Demo hardening','Current UI, input checks, deterministic contract, map, report, trace and deployment.','All mandatory flows clickable and reproducible.'),('1 / Benchmark baseline','Adapt one remote-sensing VLM; add RSVQA/VRSBench adapters and grounding.','Held-out scores, latency and failure cases recorded.'),('2 / Change intelligence','CDVQA/change captioning, registration and mask post-processing.','Temporal answers include direction, geometry and calibrated uncertainty.'),('3 / Optical–SAR ledger','Modality-specific encoders and disagreement/quality heads.','Fused result beats optical-only and SAR-only in ablation.'),('4 / Analyst workspace','AOI/CRS handling, annotation review, report templates and feedback.','Analyst can reproduce and correct an answer.'),('5 / Commercial platform','Async jobs, storage, GPU workers, tenant isolation, registry and audit log.','Repeatable SLA, cost/km², access controls and monitoring.')], [1900,5000,2960])
doc.add_heading('Startup wedge', 2)
p(doc, 'The first customer should be a high-cost decision where evidence and uncertainty matter: flood triage for insurers/disaster teams, infrastructure change monitoring for public works, or water/urban expansion intelligence for planners. The moat is the labelled evidence trail and sensor-disagreement data collected from analyst corrections, not a generic chat experience.')

doc.add_heading('10. Evaluation and acceptance criteria', 1)
table(doc, ['Dimension','Measure'], [('Answer quality','VRSBench/RSVQA/CDVQA metrics on prescribed splits with confidence intervals.'),('Spatial quality','Grounding IoU/mask IoU, CRS correctness and valid geometry rate.'),('Fusion value','Optical-only vs SAR-only vs fused ablation and sensor contribution consistency.'),('Calibration','Expected calibration error, abstention precision and error rate at release thresholds.'),('System quality','Input rejection accuracy, p50/p95 latency, trace completeness and report reproducibility.'),('Human factors','Time to first useful answer, analyst correction rate and accepted-answer percentage.')], [2400,6960])
p(doc, 'A strong demo shows one correct answer, one spatial result, one cross-modal comparison and one deliberate abstention. Showing only confident successes hides the main operational risk.')

doc.add_heading('11. Deployment and GitHub', 1)
doc.add_heading('Current web deployment', 2)
p(doc, 'The prototype is deployed as an owner-only Sites application at https://satquery-sentry.shrutikaverma22.chatgpt.site. It is suitable for review and demonstration, not public production traffic or sensitive imagery.')
doc.add_heading('Rebuild and deploy after a source change', 2)
code(doc, 'Set-Location D:\\satquery-ai\nnpm run lint\nnpm run build\n# Package dist with the Sites hosting workflow, save a version, deploy privately,\n# and poll until READY.')
p(doc, 'Source control and hosting are separate concerns: commit source, build, package deployment output, save a version, deploy, and record the deployed URL/version in release notes.')
doc.add_heading('GitHub publishing', 2)
p(doc, 'The GitHub CLI is installed but this environment is not authenticated. Run this once in PowerShell and complete the browser login:')
code(doc, 'gh auth login\n# Choose GitHub.com → HTTPS → Login with a web browser')
p(doc, 'After authentication, a sensible default for an unfinished hackathon system is a private repository:')
code(doc, 'Set-Location D:\\satquery-ai\ngh repo create satquery-ai --private --source . --remote origin --push')
p(doc, 'Use --public only after reviewing source, model weights, dataset licences and secrets. Never commit API keys, service tokens, evaluation annotations or private imagery.')

doc.add_heading('12. Troubleshooting', 1)
table(doc, ['Symptom','Likely cause','Fix'], [('Run button input error','Unsupported extension or invalid signature.','Use a real TIFF/GeoTIFF/PNG/JPEG and inspect file name/size.'),('Only curated metadata','No file uploaded for selected case.','Click ADD OBSERVATION and load one or two files.'),('No change regions','Confidence below release threshold.','Read limit/next observation; do not force a map.'),('Unexpected task','Query lacks routing cue or input count is incompatible.','Use explicit words: changed, highlight, SAR, optical, together.'),('Build failure','Node/npm or lockfile drift.','Use Node 22+, npm install, then lint and build.'),('GitHub denied','CLI not authenticated or name taken.','Run gh auth status, authenticate, choose unique name.')], [2500,4200,4060])
doc.add_heading('13. Glossary', 1)
table(doc, ['Term','Meaning'], [('AOI','Area of interest; geographic footprint being analysed.'),('SAR','Synthetic aperture radar; structural, day/night and cloud-penetrating modality.'),('Optical / multispectral','Reflectance imagery with spectral and contextual information.'),('Co-registered','Images aligned so corresponding pixels refer to the same ground location.'),('VQA','Visual question answering.'),('Grounding','Linking a phrase to a box or polygon in an image.'),('Change VQA','Answering a question about differences between corresponding observations.'),('Abstention','Withholding a conclusion when evidence is insufficient or contradictory.'),('Evidence Contract','Structured answer: claim, where, magnitude, sensor case, limit, confidence and trace.')], [2500,8260])
doc.add_heading('Appendix A. Three-minute judging script', 1)
for t in ['Open the URL and point out cases plus the evidence contract.','Run the default Assam flood query. Explain claim, location, magnitude, sensor support and limitation.','Switch OPTICAL and SAR to show modality views.','Select CASE_002 and run a single-image VQA question.','Run “Highlight water-covered regions” for spatial grounding.','Add “cloud uncertainty”; show confidence drop, withheld regions and next observation.','Expand the trace and download the report. Explain that the deterministic contract is ready for benchmarked models in the registry.']: numbered(doc,t)

OUT.parent.mkdir(parents=True, exist_ok=True); doc.save(OUT); print(OUT)
